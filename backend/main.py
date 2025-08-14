from fastapi import FastAPI, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sqlite3
import requests
import os
import datetime
from pathlib import Path
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

app = FastAPI()

# CORS setup for frontend-backend communication
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)



# Modal integration placeholder for IBM granite-3.0-1b-a4000-instruct
# TODO: Replace with correct Modal API usage for local inference

def is_finance_related(query: str) -> bool:
    """Check if query is related to finance, tax, savings, loans, or financial laws"""
    finance_keywords = [
        # Core financial terms
        'finance', 'financial', 'money', 'budget', 'budgeting', 'income', 'expense', 'expenses',
        'savings', 'save', 'saving', 'investment', 'invest', 'investing', 'portfolio',
        'loan', 'loans', 'credit', 'debt', 'mortgage', 'interest', 'rate', 'rates',
        
        # Tax related
        'tax', 'taxes', 'taxation', 'deduction', 'deductions', 'refund', 'irs',
        'filing', 'return', 'exemption', 'taxable', 'income tax', 'gst', 'vat',
        
        # Banking and accounts
        'bank', 'banking', 'account', 'checking', 'deposit', 'withdrawal',
        'atm', 'card', 'payment', 'transaction', 'balance', 'statement',
        
        # Insurance and financial products
        'insurance', 'policy', 'premium', 'claim', 'retirement', 'pension',
        '401k', 'ira', 'mutual fund', 'etf', 'stock', 'stocks', 'bond', 'bonds',
        
        # Financial planning
        'goal', 'goals', 'planning', 'wealth', 'asset', 'assets', 'liability',
        'net worth', 'cash flow', 'emergency fund', 'financial plan',
        
        # Business finance
        'business', 'startup', 'revenue', 'profit', 'loss', 'accounting',
        'bookkeeping', 'invoice', 'payroll', 'entrepreneur',
        
        # Economic terms
        'economy', 'economic', 'inflation', 'recession', 'market', 'currency',
        'exchange', 'forex', 'commodity', 'real estate', 'property'
    ]
    
    query_lower = query.lower()
    return any(keyword in query_lower for keyword in finance_keywords)

def detect_report_request(query: str) -> str:
    """Detect if user is asking for a report and return report type"""
    query_lower = query.lower()
    
    report_keywords = {
        'comprehensive': ['generate report', 'create report', 'full report', 'comprehensive report', 'detailed report'],
        'summary': ['summarize', 'summary', 'overview', 'brief report'],
        'analysis': ['analyze', 'analysis', 'assess', 'evaluate', 'review']
    }
    
    for report_type, keywords in report_keywords.items():
        if any(keyword in query_lower for keyword in keywords):
            return report_type
    return None

def generate_groq_financial_report(user_type: str, income: float, expenses: float, goal: str, goal_amount: float, chat_history: str) -> str:
    """Generate comprehensive financial report using Groq API as fallback"""
    
    # Calculate key metrics
    savings = income - expenses
    savings_rate = (savings / income * 100) if income > 0 else 0
    
    # Create detailed financial analysis prompt
    analysis_prompt = f"""
As a professional financial advisor, create a comprehensive financial report for a {user_type} with the following data:

FINANCIAL SNAPSHOT:
- Monthly Income: ${income:,.2f}
- Monthly Expenses: ${expenses:,.2f}  
- Monthly Savings: ${savings:,.2f}
- Savings Rate: {savings_rate:.1f}%
- Financial Goal: {goal}
- Goal Amount: ${goal_amount:,.2f}

CHAT HISTORY: {chat_history[:500]}...

Please provide a detailed analysis covering:
1. Current Financial Health Assessment
2. Savings and Spending Analysis
3. Goal Achievement Strategy
4. Risk Assessment and Recommendations
5. Actionable Next Steps

Keep the report professional, specific, and actionable with concrete numbers and percentages.
    """.strip()
    
    try:
        API_URL = "https://api.groq.com/openai/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {os.getenv('GROQ_API_KEY', 'your-groq-api-key-here')}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "llama3-8b-8192",
            "messages": [
                {"role": "system", "content": "You are an expert financial advisor specializing in comprehensive financial analysis and reporting. Provide detailed, specific, and actionable financial advice with concrete numbers and strategies."},
                {"role": "user", "content": analysis_prompt}
            ],
            "max_tokens": 1500,
            "temperature": 0.3  # Lower temperature for more consistent reports
        }
        
        response = requests.post(API_URL, headers=headers, json=payload, timeout=30)
        response.raise_for_status()
        result = response.json()
        
        if "choices" in result and len(result["choices"]) > 0:
            return result["choices"][0]["message"]["content"]
        else:
            return generate_structured_fallback_report(user_type, income, expenses, goal, goal_amount, savings)
            
    except Exception as e:
        return generate_structured_fallback_report(user_type, income, expenses, goal, goal_amount, savings)

def generate_structured_fallback_report(user_type: str, income: float, expenses: float, goal: str, goal_amount: float, savings: float) -> str:
    """Generate structured report when all AI services fail"""
    
    savings_rate = (savings / income * 100) if income > 0 else 0
    months_to_goal = (goal_amount / savings) if savings > 0 else float('inf')
    
    report = f"""
📊 COMPREHENSIVE FINANCIAL REPORT - {user_type.upper()}

💰 FINANCIAL SNAPSHOT
Current Monthly Income: ${income:,.2f}
Current Monthly Expenses: ${expenses:,.2f}
Net Monthly Savings: ${savings:,.2f}
Personal Savings Rate: {savings_rate:.1f}%

🎯 GOAL ANALYSIS
Target: {goal}
Required Amount: ${goal_amount:,.2f}
Time to Goal: {months_to_goal:.1f} months (at current savings rate)

📈 FINANCIAL HEALTH ASSESSMENT
{"🟢 EXCELLENT" if savings_rate >= 20 else "🟡 GOOD" if savings_rate >= 10 else "🔴 NEEDS IMPROVEMENT"}
- Your savings rate of {savings_rate:.1f}% is {"above average" if savings_rate >= 15 else "below recommended 20%"}
- {"You're building wealth effectively" if savings > 0 else "Consider reducing expenses to increase savings"}

💡 KEY RECOMMENDATIONS
1. {"Continue current savings pattern" if savings > 0 else "Create a monthly budget to identify savings opportunities"}
2. {"Consider emergency fund (3-6 months expenses)" if savings > 0 else "Focus on positive cash flow first"}
3. {"Explore investment options for goal acceleration" if savings_rate > 15 else "Optimize expenses before investing"}

📋 ACTION PLAN
□ Track expenses for 30 days
□ {"Increase savings by $" + f"{max(income*0.1 - savings, 0):.0f}" if savings < income*0.1 else "Maintain current savings level"}
□ Review and adjust monthly budget
□ {"Research investment options" if savings_rate > 10 else "Focus on debt reduction"}

Report generated using structured analysis engine.
    """
    
    return report.strip()

def create_word_report(user_type: str, income: float, expenses: float, goal: str, goal_amount: float, report_content: str, model_name: str) -> str:
    """Create a Word document from the financial report"""
    
    # Create reports directory if it doesn't exist
    reports_dir = Path("reports")
    reports_dir.mkdir(exist_ok=True)
    
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Comprehensive Financial Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add header information
    doc.add_heading('Report Details', level=1)
    details_table = doc.add_table(rows=6, cols=2)
    details_table.style = 'Table Grid'
    
    # Populate the details table
    details_data = [
        ('User Type:', user_type.title()),
        ('Generated On:', datetime.datetime.now().strftime('%B %d, %Y at %I:%M %p')),
        ('Monthly Income:', f'${income:,.2f}'),
        ('Monthly Expenses:', f'${expenses:,.2f}'),
        ('Financial Goal:', goal if goal else 'Not specified'),
        ('Goal Amount:', f'${goal_amount:,.2f}' if goal_amount > 0 else 'Not specified')
    ]
    
    for i, (label, value) in enumerate(details_data):
        row = details_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
    
    # Add spacing
    doc.add_paragraph()
    
    # Add AI analysis section
    doc.add_heading('AI Financial Analysis', level=1)
    
    # Clean and format the report content
    report_lines = report_content.split('\n')
    current_paragraph = None
    
    for line in report_lines:
        line = line.strip()
        if not line:
            continue
            
        # Handle headings (lines with emojis or all caps)
        if any(emoji in line for emoji in ['📊', '💰', '🎯', '📈', '💡', '📋', '🟢', '🟡', '🔴', '⚡', '⚠️']):
            if line.startswith('📊') or 'REPORT' in line.upper():
                doc.add_heading(line.replace('📊', '').strip(), level=2)
            elif any(emoji in line for emoji in ['💰', '🎯', '📈', '💡', '📋']):
                doc.add_heading(line, level=3)
            else:
                doc.add_paragraph(line)
        # Handle bullet points
        elif line.startswith(('□', '-', '•', '1.', '2.', '3.', '4.', '5.')):
            doc.add_paragraph(line, style='List Bullet')
        # Handle regular paragraphs
        else:
            doc.add_paragraph(line)
    
    # Add footer with model information
    doc.add_paragraph()
    doc.add_heading('Report Generation Information', level=1)
    footer_para = doc.add_paragraph()
    footer_para.add_run(f'This report was generated using: {model_name}').bold = True
    footer_para.add_run(f'\nGenerated on: {datetime.datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
    footer_para.add_run('\n\nThis report is for informational purposes only and should not be considered as professional financial advice.')
    
    # Generate filename
    timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"financial_report_{user_type}_{timestamp}.docx"
    filepath = reports_dir / filename
    
    # Save the document
    doc.save(str(filepath))
    
    return str(filepath)

def run_model(prompt: str, user_type: str) -> str:
    # Check if query is finance-related
    if not is_finance_related(prompt):
        return "I'm a specialized financial assistant. I can only help with questions related to finance, tax, savings, loans, investments, budgeting, and financial planning. Please ask me about financial topics!"
    
    # Check if user is requesting a report
    report_type = detect_report_request(prompt)
    if report_type:
        return f"I can help you generate a {report_type} financial report! Please use the 'Generate Report' button in the interface, or provide your financial data (income, expenses, goals) and I'll create a comprehensive analysis using IBM Granite AI."
    
    # Groq API integration
    API_URL = "https://api.groq.com/openai/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {os.getenv('GROQ_API_KEY', 'your-groq-api-key-here')}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": "llama3-8b-8192",
        "messages": [
            {"role": "system", "content": "You are a specialized financial assistant. You ONLY answer questions related to finance, tax, savings, loans, investments, budgeting, financial planning, banking, insurance, and financial laws. If asked about anything else, politely redirect the conversation back to financial topics. Provide helpful, accurate financial advice based on the user's profile."},
            {"role": "user", "content": f"As a {user_type}, {prompt}"}
        ],
        "max_tokens": 256,
        "temperature": 0.7
    }
    try:
        response = requests.post(API_URL, headers=headers, json=payload, timeout=30)
        response.raise_for_status()
        result = response.json()
        if "choices" in result and len(result["choices"]) > 0:
            return result["choices"][0]["message"]["content"]
        else:
            return str(result)
    except requests.exceptions.HTTPError as e:
        try:
            err_json = e.response.json()
            return f"AI model error: {err_json.get('error', str(e))}"
        except Exception:
            return f"AI model error: {e}"
    except Exception as e:
        return f"AI model error: {e}"

@app.post("/chat")
async def chat(request: Request):
    data = await request.json()
    user_message = data.get("message", "")
    user_type = data.get("user_type", "student")
    language = data.get("language", "english")
    # Supported Indian languages
    indian_languages = {
        "hindi": "Hindi",
        "bengali": "Bengali",
        "tamil": "Tamil",
        "telugu": "Telugu",
        "marathi": "Marathi",
        "gujarati": "Gujarati",
        "kannada": "Kannada",
        "malayalam": "Malayalam",
        "punjabi": "Punjabi"
    }
    if language.lower() in indian_languages:
        prompt = f"Reply in {indian_languages[language.lower()]}: {user_message}"
    else:
        prompt = user_message
    response = run_model(prompt, user_type)
    return {"response": response}

@app.post("/budget-summary")
async def budget_summary(request: Request):
    data = await request.json()
    income = data.get("income", 0)
    expenses = data.get("expenses", 0)
    savings = income - expenses
    if savings < 0:
        summary = f"Your expenses exceed your income by ${-savings:.2f}. Consider reducing discretionary spending."
    elif savings == 0:
        summary = "You are breaking even. Try to save a small amount each month for emergencies."
    else:
        summary = f"You are saving ${savings:.2f} per month. Great job! Consider allocating some to your financial goals."
    return {"summary": summary}

@app.post("/goal-calculation")
async def goal_calculation(request: Request):
    data = await request.json()
    goal_amount = data.get("goal_amount", 0)
    income = data.get("income", 0)
    expenses = data.get("expenses", 0)
    months = data.get("months", 12)  # Default to 12 months if not provided
    available = max(income - expenses, 0)
    if months > 0:
        monthly_needed = goal_amount / months
    else:
        monthly_needed = goal_amount
    if available < monthly_needed:
        msg = f"You need to save ${monthly_needed:.2f} per month, but only ${available:.2f} is available. Consider extending your timeline or reducing expenses."
    else:
        msg = f"You need to save ${monthly_needed:.2f} per month to reach your goal. You have enough available!"
    return {"monthly_amount": monthly_needed, "message": msg}

@app.post("/spending-insights")
async def spending_insights(request: Request):
    data = await request.json()
    expenses = data.get("expenses", 0)
    # For demo, highlight if expenses are high compared to income
    income = data.get("income", 0)
    if income > 0 and expenses/income > 0.7:
        insights = "Your spending is more than 70% of your income. Review subscriptions, eating out, and impulse purchases."
    else:
        insights = "Your spending is within a healthy range. Keep tracking for overlooked expenses like small subscriptions or fees."
    return {"insights": insights}

@app.post("/save-session")
async def save_session(request: Request):
    data = await request.json()
    user_type = data.get("user_type", "student")
    chat_history = data.get("chat_history", "")
    income = data.get("income", 0)
    expenses = data.get("expenses", 0)
    goal = data.get("goal", "")
    goal_amount = data.get("goal_amount", 0)
    try:
        conn = sqlite3.connect("financebot.db")
        c = conn.cursor()
        c.execute("""
            INSERT INTO sessions (user_type, chat_history, income, expenses, goal, goal_amount)
            VALUES (?, ?, ?, ?, ?, ?)
        """, (user_type, chat_history, income, expenses, goal, goal_amount))
        conn.commit()
        session_id = c.lastrowid
        conn.close()
        return {"status": "Session saved.", "session_id": session_id}
    except Exception as e:
        return {"status": f"Error saving session: {e}"}
@app.get("/get-session/{session_id}")
async def get_session(session_id: int):
    try:
        conn = sqlite3.connect("financebot.db")
        c = conn.cursor()
        c.execute("SELECT * FROM sessions WHERE id = ?", (session_id,))
        row = c.fetchone()
        conn.close()
        if row:
            keys = ["id", "user_type", "chat_history", "income", "expenses", "goal", "goal_amount", "created_at"]
            session = dict(zip(keys, row))
            return {"session": session}
        else:
            return {"error": "Session not found."}
    except Exception as e:
        return {"error": str(e)}

@app.post("/generate-comprehensive-report")
async def generate_comprehensive_report(request: Request):
    """
    Generate comprehensive financial report using IBM Granite model
    """
    data = await request.json()
    user_type = data.get("user_type", "student")
    income = data.get("income", 0)
    expenses = data.get("expenses", 0)
    goal = data.get("goal", "")
    goal_amount = data.get("goal_amount", 0)
    chat_history = data.get("chat_history", "")
    
    # Prepare raw content for Granite model
    raw_content = f"""
user_type: {user_type}
income: {income}
expenses: {expenses}
goal: {goal}
goal_amount: {goal_amount}
chat_history: {chat_history}
    """.strip()
    
    try:
        # Try Granite report service first with shorter timeout
        granite_response = requests.post(
            "http://localhost:8002/generate-report",
            json={
                "raw_content": raw_content,
                "report_type": "comprehensive_financial_analysis"
            },
            timeout=90  # Reduced timeout for faster fallback
        )
        
        if granite_response.status_code == 200:
            report_data = granite_response.json()
            return {
                "report": report_data.get("report", "Report generation failed"),
                "model": "IBM Granite 3.0-1B",
                "status": "success"
            }
        else:
            # Fall back to Groq-powered structured report
            fallback_report = generate_groq_financial_report(user_type, income, expenses, goal, goal_amount, chat_history)
            return {
                "report": fallback_report,
                "model": "Groq API (llama3-8b-8192) - Structured Analysis",
                "status": "success"
            }
            
    except requests.exceptions.ReadTimeout:
        # Granite model is taking too long, use Groq fallback
        fallback_report = generate_groq_financial_report(user_type, income, expenses, goal, goal_amount, chat_history)
        return {
            "report": fallback_report + "\n\n⚡ Note: Generated using fast AI analysis due to high demand on specialized models.",
            "model": "Groq API (llama3-8b-8192) - Fast Analysis",
            "status": "success"
        }
    except Exception as e:
        # Any other error, use Groq fallback
        fallback_report = generate_groq_financial_report(user_type, income, expenses, goal, goal_amount, chat_history)
        return {
            "report": fallback_report + f"\n\n⚠️ Note: Fallback analysis used due to: {str(e)[:100]}",
            "model": "Groq API (llama3-8b-8192) - Fallback Analysis", 
            "status": "success"
        }

@app.post("/analyze-my-finances")
async def analyze_my_finances(request: Request):
    """
    Analyze user's financial situation using session data and Granite AI
    """
    data = await request.json()
    session_id = data.get("session_id")
    
    if not session_id:
        return {"error": "Session ID required for financial analysis"}
    
    try:
        # Call the Granite analysis service
        granite_response = requests.post(
            "http://localhost:8002/analyze-session",
            json={"session_id": session_id},
            timeout=180
        )
        
        if granite_response.status_code == 200:
            analysis_data = granite_response.json()
            return {
                "analysis": analysis_data.get("report", "Analysis failed"),
                "session_data": analysis_data.get("session_data", {}),
                "model": "IBM Granite 3.0-1B",
                "status": "success"
            }
        else:
            return {"error": "Granite analysis service unavailable", "status": "error"}
            
    except Exception as e:
        return {"error": f"Financial analysis failed: {str(e)}", "status": "error"}

@app.post("/generate-word-report")
async def generate_word_report(request: Request):
    """
    Generate comprehensive financial report as downloadable Word document
    """
    data = await request.json()
    user_type = data.get("user_type", "student")
    income = data.get("income", 0)
    expenses = data.get("expenses", 0)
    goal = data.get("goal", "")
    goal_amount = data.get("goal_amount", 0)
    chat_history = data.get("chat_history", "")
    
    # First generate the report content using the same logic as the regular report
    raw_content = f"""
user_type: {user_type}
income: {income}
expenses: {expenses}
goal: {goal}
goal_amount: {goal_amount}
chat_history: {chat_history}
    """.strip()
    
    try:
        # Try Granite report service first with shorter timeout
        granite_response = requests.post(
            "http://localhost:8002/generate-report",
            json={
                "raw_content": raw_content,
                "report_type": "comprehensive_financial_analysis"
            },
            timeout=90  # Reduced timeout for faster fallback
        )
        
        if granite_response.status_code == 200:
            report_data = granite_response.json()
            report_content = report_data.get("report", "Report generation failed")
            model_name = "IBM Granite 3.0-1B"
        else:
            # Fall back to Groq-powered structured report
            report_content = generate_groq_financial_report(user_type, income, expenses, goal, goal_amount, chat_history)
            model_name = "Groq API (llama3-8b-8192) - Structured Analysis"
            
    except requests.exceptions.ReadTimeout:
        # Granite model is taking too long, use Groq fallback
        report_content = generate_groq_financial_report(user_type, income, expenses, goal, goal_amount, chat_history)
        model_name = "Groq API (llama3-8b-8192) - Fast Analysis"
    except Exception as e:
        # Any other error, use Groq fallback
        report_content = generate_groq_financial_report(user_type, income, expenses, goal, goal_amount, chat_history)
        model_name = "Groq API (llama3-8b-8192) - Fallback Analysis"
    
    # Create Word document
    try:
        word_file_path = create_word_report(user_type, income, expenses, goal, goal_amount, report_content, model_name)
        
        # Return file for download
        return FileResponse(
            path=word_file_path,
            filename=os.path.basename(word_file_path),
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return {"error": f"Word document generation failed: {str(e)}", "status": "error"}

@app.get("/health")
async def health():
    return {"status": "ok"}
